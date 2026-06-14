"""
Markdown -> well-formatted .docx converter for the Agentic Hyper Scalers specs.

Handles the Markdown subset used in our docs:
  # / ## / ### / ####  headings        | tables (with header shading)
  - / nested bullets    1. numbered     inline **bold** *italic* `code` [text](url)
  > blockquotes         --- rules       ``` code / ```mermaid (rendered as code block w/ caption)

Usage:  python md_to_docx.py <input.md> <output.docx> "<Doc Title>" "<Doc Subtitle>"
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x0B, 0x66, 0x8C)      # deep cyan for headings
ACCENT_BG = "11638C"                      # header-row shading
CODE_BG = "F2F4F7"
LINK_BLUE = RGBColor(0x0B, 0x57, 0xD0)
CODE_RED = RGBColor(0x9B, 0x2C, 0x2C)

INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|\*[^*]+\*)')


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def shade_paragraph(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C7CDD6')
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0B57D0'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_runs(paragraph, text):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`'):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = 'Consolas'; r.font.size = Pt(9.5); r.font.color.rgb = CODE_RED
        elif tok.startswith('['):
            mm = re.match(r'\[([^\]]+)\]\(([^)]+)\)', tok)
            add_hyperlink(paragraph, mm.group(2), mm.group(1))
        elif tok.startswith('*'):
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def style_base(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    for i, sz in [(1, 17), (2, 13.5), (3, 11.5), (4, 10.5)]:
        st = doc.styles[f'Heading {i}']
        st.font.name = 'Calibri'
        st.font.size = Pt(sz)
        st.font.color.rgb = ACCENT
        st.font.bold = True


def add_footer(doc, label):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + "   |   Page ")
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x6B, 0x76, 0x85)
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    r = p.add_run(); r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x6B, 0x76, 0x85)
    r._r.append(fld_begin); r._r.append(instr); r._r.append(fld_end)


def add_table(doc, rows):
    header = rows[0]
    header_empty = all(c.strip() == '' for c in header)
    body = rows[1:] if header_empty else rows
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ridx, row in enumerate(body):
        cells = table.add_row().cells
        is_header = (not header_empty) and ridx == 0
        for cidx in range(ncols):
            txt = row[cidx] if cidx < len(row) else ''
            cell = cells[cidx]
            cell.paragraphs[0].text = ''
            p = cell.paragraphs[0]
            if is_header:
                set_cell_bg(cell, ACCENT_BG)
                r = p.add_run(txt.replace('**', ''))
                r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(9.5)
            else:
                add_runs(p, txt)
                for rr in p.runs:
                    rr.font.size = Pt(9.5)
    doc.add_paragraph()


def add_code_block(doc, lines, lang):
    if lang == 'mermaid':
        cap = doc.add_paragraph()
        cr = cap.add_run('Figure — architecture diagram (Mermaid source; renders graphically on GitHub):')
        cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = RGBColor(0x6B, 0x76, 0x85)
    p = doc.add_paragraph()
    shade_paragraph(p, CODE_BG)
    run = p.add_run('\n'.join(lines))
    run.font.name = 'Consolas'; run.font.size = Pt(8.5); run.font.color.rgb = RGBColor(0x24, 0x2B, 0x33)
    doc.add_paragraph()


def parse_table_row(line):
    s = line.strip()
    if s.startswith('|'):
        s = s[1:]
    if s.endswith('|'):
        s = s[:-1]
    return [c.strip() for c in s.split('|')]


def is_separator(line):
    return bool(re.match(r'^\s*\|?\s*:?-{2,}.*\|', line)) and set(line.strip()) <= set('|-: ')


def convert(md_path, docx_path, title, subtitle, footer_label):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().replace('\r\n', '\n').split('\n')

    doc = Document()
    style_base(doc)
    add_footer(doc, footer_label)

    # ---- Cover ----
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = t.add_run(title); tr.bold = True; tr.font.size = Pt(26); tr.font.color.rgb = ACCENT
    s = doc.add_paragraph(); sr = s.add_run(subtitle)
    sr.font.size = Pt(13); sr.font.color.rgb = RGBColor(0x44, 0x4D, 0x59)
    add_bottom_border(doc.add_paragraph())

    doc.core_properties.title = title
    doc.core_properties.author = 'Raghuram Nimishakavi'

    i = 0
    n = len(lines)
    first_h1_skipped = False
    skip_subtitle = False
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Skip the doc's own first H1 + the bold subtitle line (already on cover)
        if not first_h1_skipped and stripped.startswith('# '):
            first_h1_skipped = True
            skip_subtitle = True
            i += 1
            continue
        if skip_subtitle:
            skip_subtitle = False
            if stripped.startswith('**') and stripped.endswith('**'):
                i += 1
                continue

        if stripped == '':
            i += 1
            continue

        # Code fence
        if stripped.startswith('```'):
            lang = stripped[3:].strip().lower()
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                block.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, block, lang)
            continue

        # Table
        if stripped.startswith('|') and i + 1 < n and is_separator(lines[i + 1]):
            rows = [parse_table_row(line)]
            i += 1  # skip header
            i += 1  # skip separator
            while i < n and lines[i].strip().startswith('|'):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        # Horizontal rule
        if re.match(r'^\s*---+\s*$', line):
            add_bottom_border(doc.add_paragraph())
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = min(len(m.group(1)), 4)
            p = doc.add_heading(level=level)
            add_runs(p, m.group(2))
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, stripped.lstrip('> ').rstrip())
            for r in p.runs:
                r.italic = True; r.font.color.rgb = RGBColor(0x44, 0x4D, 0x59)
            i += 1
            continue

        # Numbered list
        mnum = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
        if mnum:
            indent = len(mnum.group(1))
            style = 'List Number' if indent < 2 else 'List Number 2'
            p = doc.add_paragraph(style=style)
            add_runs(p, mnum.group(3))
            i += 1
            continue

        # Bullet list (supports nesting by indent)
        mbul = re.match(r'^(\s*)[-*]\s+(.*)$', line)
        if mbul:
            indent = len(mbul.group(1))
            lvl = min(indent // 2, 2)
            style = 'List Bullet' if lvl == 0 else f'List Bullet {lvl + 1}'
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style='List Bullet')
            add_runs(p, mbul.group(2))
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f'Wrote {docx_path}')


if __name__ == '__main__':
    convert(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4], sys.argv[5])
